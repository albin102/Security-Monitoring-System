import cv2
import numpy as np
from ultralytics import YOLO
from keras.models import load_model
from mtcnn.mtcnn import MTCNN
from keras_facenet import FaceNet
from keras.applications.inception_resnet_v2 import preprocess_input
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import docx
from docx.shared import Inches
import datetime
from geopy.geocoders import Nominatim

# Email credentials
email_sender = "albinkuriachan24@gmail.com"
email_receiver = "akuriachan39@gmail.com"
sender_password = "tkij crew dvaz hjou"

# Initialize models
wp_model = YOLO("best.pt")
facenet = FaceNet()
model_save_path = 'Suspicious Detection/Model/'
model = load_model(os.path.join(model_save_path, 'neural_network_model12.h5'))
face_detector = MTCNN()

def get_current_location():
    loc = Nominatim(user_agent="GetLoc")
    getLoc = loc.geocode("Newport, Wales, United Kingdom")
    if getLoc:
        return getLoc.address
    else:
        return "Location not found"

def send_email_with_attachment(email_sender, email_receiver, sender_password, attachment_path, subject):
    msg = MIMEMultipart()
    msg['From'] = email_sender
    msg['To'] = email_receiver
    msg['Subject'] = subject

    # Attach file
    with open(attachment_path, "rb") as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
    encoders.encode_base64(part)
    part.add_header(
        "Content-Disposition",
        f"attachment; filename= {attachment_path}",
    )
    msg.attach(part)

    # Connect to SMTP server and send email
    with smtplib.SMTP("smtp.gmail.com", 587) as server:
        server.starttls()
        server.login(email_sender, sender_password)
        server.sendmail(email_sender, email_receiver, msg.as_string())

def wpreport():
    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    image_path = 'weapon_frame.jpg'
    current_location = get_current_location()
    dcm = "Weapon Detected"
    doc = docx.Document()
    doc.add_heading('Report', 0)
    para = doc.add_paragraph(dcm)
    para.paragraph_format.line_spacing = Inches(0.5)
    doc.add_picture(image_path, width=Inches(4))
    doc.add_paragraph(f"Current Location: {current_location}")
    doc.add_paragraph(f"Current Time: {current_time}")
    doc.save("Weapon_report.docx")

def weapon_detection(frame):
    results_w = wp_model(frame, show=True)
    print(results_w)
    confidence_w = np.array(results_w[0].boxes.conf)

    print("####################")
    print("#Weapon-Detection#")
    print("####################")
    print(results_w)
    print("####################")
    print("####################")
    print("********************")
    if np.any((confidence_w >= 0.80)):  
        print("Weapon detected")
        cv2.imwrite(f"weapon_frame.jpg", frame)
        wpreport()
        send_email_with_attachment(email_sender, email_receiver, sender_password, "Weapon_report.docx", "Weapon Detected!")

def crreport(nm):
    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    image_path = 'criminal_detected_frame.jpg'
    current_location = get_current_location()
    dcm = "Criminal Detected: " + nm
    doc = docx.Document()
    doc.add_heading('Report', 0)
    para = doc.add_paragraph(dcm)
    para.paragraph_format.line_spacing = Inches(0.5)
    doc.add_picture(image_path, width=Inches(4))
    doc.add_paragraph(f"Current Location: {current_location}")
    doc.add_paragraph(f"Current Time: {current_time}")
    doc.save("criminal_report.docx")

def face_det(frame):
    class_idx_to_celebrity = {0: 'criminal1', 1: 'criminal2'}
    resized_width = 800
    resized_height = 600

    frame = cv2.resize(frame, (resized_width, resized_height))
    rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    faces = face_detector.detect_faces(rgb_frame)

    for face in faces:
        x, y, w, h = face['box']
        face_img = frame[y:y+h, x:x+w]
        face_img = cv2.resize(face_img, (160, 160))
        face_img = np.expand_dims(face_img, axis=0)
        embedding = facenet.model.predict(preprocess_input(face_img))[0]
        predictions = model.predict(np.array([embedding]))
        predicted_class_idx = np.argmax(predictions)
        confidence = predictions[0][predicted_class_idx]
        threshold = 0.95

        if confidence > threshold:
            celebrity_name = class_idx_to_celebrity.get(predicted_class_idx, 'Unknown')
            color = (0, 255, 0)
            text = f"{celebrity_name} ({confidence:.2f})"
            cv2.imwrite(f"criminal_detected_frame.jpg", frame)
            if celebrity_name == "criminal2":
                celebrity_name = "Ranbir Kapoor"
            else:
                celebrity_name = "Irfan Khan"
            crreport(celebrity_name)
            res = "Suspicious Person: " + celebrity_name
            send_email_with_attachment(email_sender, email_receiver, sender_password, "criminal_report.docx", res)
        else:
            color = (0, 0, 255)
            text = "Unknown"

        cv2.rectangle(frame, (x, y), (x+w, y+h), color, 2)
        cv2.putText(frame, text, (x, y + 20), cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 2)

    ff = cv2.resize(frame, (500, 500))
    cv2.imshow('Face', ff)
    cv2.waitKey(1)

def image_check(image_path):
    frame = cv2.imread(image_path)
    if frame is not None:
        weapon_detection(frame)
        face_det(frame)
    else:
        print("Failed to load image.")